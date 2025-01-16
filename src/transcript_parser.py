import json
import re
from pathlib import Path
from typing import List

import pdfplumber
from docx import Document
from loguru import logger
from tqdm import tqdm

from src.classes import TranscriptClass

MATCH_TEXT_START_LENGTH = 20
MATCH_TEXT_END_LENGTH = -20


class TranscriptionParser:
    def __init__(self, dir_path: Path, output_path: Path):
        self.dir_path = dir_path
        self.output_path = output_path

    def _fulltext_with_page_number(self, filename: Path):
        result = []
        with pdfplumber.open(filename) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                page_number = text.split("\n")[-1]
                if not any(char.isdigit() for char in page_number):
                    continue
                page_number = page_number
                text = self._remove_spaces(text)
                result.append((page_number, text))
        return result

    def _remove_spaces(self, s: str):
        s = s.replace("\n", "")
        s = re.sub(r"(?<=[\u4e00-\u9fff]) (?=[\u4e00-\u9fff0-9])", "", s)
        s = re.sub(r"(?<=[\u4e00-\u9fff0-9]) (?=[\u4e00-\u9fff])", "", s)
        return s

    def _get_page_number(self, text: str, fulltext_with_page_number: List[str]):
        if "\n" in text:
            text = text.split("\n")[0]
        text = self._remove_spaces(text)
        for page_number, fulltext in fulltext_with_page_number:
            if text in fulltext:
                return page_number
        return None

    def parse_document(self, filename: Path) -> List[TranscriptClass]:
        fulltext_with_page_number = self._fulltext_with_page_number(filename.with_suffix(".pdf"))

        doc = Document(filename)
        len_paragraphs = len(doc.paragraphs)

        result = []
        for i in tqdm(range(len_paragraphs), desc=f"Parsing {filename.stem}", leave=False):
            text = doc.paragraphs[i].text
            content = [x for x in doc.paragraphs[i].iter_inner_content()]

            if len(content) == 0:
                continue

            # if "：" in content[0].text and content[0].bold:
            if "：" in content[0].text[:10]:
                speaker = content[0].text
                j = i
                while True:
                    if j >= len_paragraphs - 1:
                        break
                    j += 1
                    content = [x for x in doc.paragraphs[j].iter_inner_content()]

                    if len(content) == 0:
                        continue

                    if "：" not in content[0].text and not content[0].bold:
                        text += "\n"
                        text += doc.paragraphs[j].text
                    else:
                        break
                text = text.replace(speaker, "")
                text = self._remove_spaces(text)
                speaker = speaker.replace("：", "")
                page_number_start = self._get_page_number(text[:MATCH_TEXT_START_LENGTH], fulltext_with_page_number)
                page_number_end = self._get_page_number(text[MATCH_TEXT_END_LENGTH:], fulltext_with_page_number)
                output = TranscriptClass(
                    filename=filename.stem, index=i, page_number=page_number_start if page_number_start else page_number_end, speaker=speaker, content=text
                )
                result.append(output)
        return result

    def run(self):
        logger.info("Parsing .docx files...")
        docx_file = list(self.dir_path.glob("*.docx"))
        total_files = len(docx_file)

        pbar = tqdm(docx_file, total=total_files, desc="Word to Json")

        for file in docx_file:
            pbar.set_description(f"Converting to json: {file.name}")
            save_path = self.output_path / (file.stem + ".json")

            if save_path.exists():
                pbar.update(1)
                continue

            result: List[TranscriptClass] = self.parse_document(file)

            if not result:
                pbar.update(1)
                continue

            with open(file=save_path, mode="w", encoding="utf-8") as f:
                data = json.dumps([x.model_dump() for x in result], indent=4, ensure_ascii=False)
                f.write(data)

            pbar.update(1)

        logger.info("Done parsing .docx files.")
